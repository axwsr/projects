from sqlalchemy.orm import Session
from fastapi import APIRouter,Depends,HTTPException,BackgroundTasks
from fastapi.responses import StreamingResponse
from db.session import get_session
from schemas.Document import *
from schemas.Hour_work import *
from controllers.Hour_work import *
from controllers.Person import *
from core.config import settings
from core.utils import generate_uuid, get_current_date,tuple_to_dict
import convertapi

import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

router = APIRouter()

@router.post("/generate_document")
def generate_PDF(background_tasks: BackgroundTasks,schema_doc:document_base,db:Session = Depends(get_session)):
    info_personal = get_person_information(db)
    current_date = get_current_date()
    schema_get = hour_get(start_date=schema_doc.start_date,end_date=schema_doc.end_date)
    
    hours_worked = get_hour_work_by_date(db,schema_get)
    hours_worked_dicts = [tuple_to_dict(record) for record in hours_worked]
    
    DATOS_EJEMPLO = {
        "current_date": current_date,
        "name": info_personal.first_name,
        "apellido": info_personal.last_name,
        "cursive": info_personal.first_name,
        "ojito": info_personal.last_name,
        "document": info_personal.identity_document,
        "pay_number": schema_doc.pay_number,
        "pay_word": schema_doc.pay_word,
        "observation": schema_doc.observation,
        "total_hrs": f"\n\n {schema_doc.total_hrs}",
        "address": info_personal.address_person,
        "cell_phone": info_personal.cell_phone,
        "records": hours_worked_dicts
    }
    
    template_path = "document/example.docx"
    
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    doc = Document(template_path)

    styles = doc.styles
    cursive_style = styles.add_style('CursiveMayuscula', WD_STYLE_TYPE.CHARACTER)
    cursive_style.font.italic = True
    cursive_style.font.bold = True
    cursive_style.font.all_caps = True
    cursive_style.font.size = Pt(12)

    for paragraph in doc.paragraphs:
        for key, value in DATOS_EJEMPLO.items():
            if key != "records":
                if f"{key}" in paragraph.text:
                    if key == "ojito":
                        start = paragraph.text.index(f"{key}")
                        end = start + len(f"{key}")
                        run = paragraph.runs[0]
                        run.text = run.text.replace(f"{key}", value)
                        run.style = 'CursiveMayuscula'
                    else:
                        paragraph.text = paragraph.text.replace(f"{key}", str(value))
                        
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for registro in DATOS_EJEMPLO["records"]:
            new_row = table.add_row()
            for i, value in enumerate(registro.values()):
                new_row.cells[i].text = str(value)
        
        total_row = table.add_row()
        total_row.cells[4].text = "\n\n Total Horas"
        total_row.cells[5].text = DATOS_EJEMPLO["total_hrs"]
    
    name_docx = generate_uuid()
    output_path = f"document/{name_docx}.docx"
    doc.save(output_path)
    
    convertapi.api_secret = settings.CONVERT_API_SECRET
    result = convertapi.convert('pdf', {'File': output_path})
    name_pdf = generate_uuid()
    pdf_path = f"document/{name_pdf}.pdf"
    result.file.save(pdf_path)

    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=500, detail="Error al convertir el documento a PDF")

    pdf_file = open(pdf_path, 'rb')
    name_file = f"attachment; filename={DATOS_EJEMPLO['apellido']}_horas_trabajas.pdf"
    header = {
        "Content-Disposition": name_file
    }
    
    def cleanup():
        pdf_file.close()
        os.remove(output_path)
        os.remove(pdf_path)

    background_tasks.add_task(cleanup)
    return StreamingResponse(pdf_file, media_type='application/pdf', headers=header)